import pymysql
from docxtpl import DocxTemplate
from docx import Document
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_COLOR_INDEX
from flask import Flask, render_template, request, make_response
import docx
import os
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import RGBColor
from werkzeug.utils import redirect
import datetime

date = datetime.datetime.today().strftime('%d-%m-%Y')

mydb = pymysql.connect(
  host="localhost",
  user="root",
  password="",
  db="my_python"
)
myCursor = mydb.cursor()

mydb.commit()

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    r = paragraph.add_run ()
    r._r.append (hyperlink)
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True
    return hyperlink

app = Flask(__name__)




@app.route('/')
def index():
    if request.authorization and request.authorization.username == 'admin' and request.authorization.password == 'aymen':

        return render_template('/Accueil.html')
    return make_response('Could not verify!', 401, {'WWW-Authenticate': 'Basic realm="login Required=False"'})


@app.route('/post-login')
def index1():
    return render_template('/index.html')
@app.route('/post-login', methods=['get','post'])
def login():
    sujet = request.form['sujet']
    type = request.form['type']
    serv = request.form['serv']
    numero = request.form['numero']
    source = request.form['source']
    client = request.form['client']
    description = request.form['description']

    doc = Document()
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "\tCyber Security Innovation – Response Team"
    font = paragraph.style.font
    font.name = 'Hanzel Extended'
    font.size = Pt(14)
    font.color.rgb = RGBColor(19, 32, 223)


    parag = doc.add_paragraph()

    titre1 = parag.add_run('\t \t \tSecurity Bulletin N°' + client + ' ')
    font = titre1.font
    font.name = 'Cambria (Headings)'
    font.size = Pt(14)

    titre2 = parag.add_run(type)
    font = titre2.font
    font.name = 'Cambria (Headings)'
    font.size = Pt(14)

    font.bold = True

    if (serv == "Elévé"):
        font.color.rgb = RGBColor(255, 0, 0)
    elif (serv == "Moyenne"):
        font.color.rgb = RGBColor(255, 128, 0)
    else:
        font.color.rgb = RGBColor(0, 153, 0)



    titre3 = parag.add_run('-')
    font = titre3.font
    font.name = 'Cambria (Headings)'
    font.size = Pt(14)

    font.bold = True


    if (serv == "Elévé"):
        font.color.rgb = RGBColor(255, 0, 0)
    elif (serv == "Moyenne"):
        font.color.rgb = RGBColor(255, 51, 102)
    else:
        font.color.rgb = RGBColor(0, 153, 0)

    titre4 = parag.add_run(date + '-' + numero + '\n')
    font = titre4.font
    font.name = 'Cambria (Headings)'
    font.size = Pt(14)
    font.color.rgb = RGBColor(255, 128, 0)
    font.bold = True

    if (serv == "Elévé"):
        font.color.rgb = RGBColor(255, 0, 0)
    elif (serv == "Moyenne"):
        font.color.rgb = RGBColor(255, 128, 0)
    else:
        font.color.rgb = RGBColor(0, 153, 0)

    parag.add_run('\n')

    titre5 = parag.add_run('Sujet: '+sujet+'\n')
    if " " in sujet:
        sujet = sujet.replace(" ","_")
    font = titre5.font
    font.name = 'Calibri (Body)'
    font.size = Pt(14)
    font.bold = True
    font.color.rgb = RGBColor(255, 255, 255)
    if (serv == "Elévé"):
        font.highlight_color = WD_COLOR_INDEX.RED
    elif (serv == "Moyenne"):
        font.highlight_color = WD_COLOR_INDEX.YELLOW
    else:
        font.highlight_color = WD_COLOR_INDEX.GREEN

    titre6 = parag.add_run("Référence:")
    font = titre6.font
    font.name = 'Cambria (Headings)'
    font.size = Pt(14)
    font.bold = True

    titre7 = parag.add_run(type + '-' + date + '-'+numero + '\n')
    font = titre7.font
    font.name = 'Calibri (Body)'
    font.size = Pt(16)

    font.bold = True
    if (serv == "Elévé"):
        font.color.rgb = RGBColor(255, 0, 0)
    elif (serv == "Moyenne"):
        font.color.rgb = RGBColor(255, 128, 0)
    else:
        font.color.rgb = RGBColor(0, 153, 0)

    titre8 = parag.add_run("Sévérité: ")
    font = titre8.font
    font.name = 'Calibri (Body)'
    font.size = Pt(14)
    font.bold = True

    titre9 = parag.add_run(serv + '\n')
    font = titre9.font
    font.name = 'Liberation Serif'
    font.size = Pt(16)

    font.bold = True
    if (serv == "Elévé"):
        font.color.rgb = RGBColor(255, 0, 0)
    elif (serv == "Moyenne"):
        font.color.rgb = RGBColor(255, 128, 0)
    else:
        font.color.rgb = RGBColor(0, 153, 0)

    titre10 = parag.add_run('Description: \n')
    font = titre10.font
    font.name = 'Calibri (Body)'
    font.size = Pt(14)
    font.bold = True

    titre11 = parag.add_run(description + '\n')
    font = titre11.font
    font.name = 'Calibri (Body)'
    font.size = Pt(12)

    titre12 = parag.add_run('Source: \n')
    font = titre12.font
    font.name = 'Calibri (Body)'
    font.size = Pt(14)
    font.bold = True


    p = doc.add_paragraph('')
    add_hyperlink(p, source, source)

    doc.save("Bulletin_"+client+"-"+type+"_"+date+"_"+"0"+numero+"_"+sujet+".docx")

    os.system("start Bulletin_" + client + "-" + type + "_" + date + "_" +"0" + numero + "_" + sujet + ".docx")
    return render_template('index.html')
@app.route('/rapport')
def rapport1():
    return render_template('/rapport.html')


@app.route('/rapport', methods=['get', 'post'])
def rapport():
    sujet = request.form['sujet']
    type = request.form['type']
    serv = request.form['serv']
    numero = request.form['numero']
    img = request.form['description']
    description = request.form['description']
    logmesg = request.form['logmesg']
    recommandation = request.form['recommandation']

    requete = """INSERT INTO atb(type,numero,date,sujet,description,log_mesg,recommandation) VALUES(%s, %s, %s, %s, %s, %s)"""
    recordTuple = (type, numero, date, sujet, description, logmesg, recommandation)
    myCursor.execute(requete, recordTuple)
    mydb.commit()

    doc = DocxTemplate("ATB_Template.docx")
    context = {'Keystone': "World company"}
    doc.render(context)
    parag = doc.add_paragraph()

    titre1 = parag.add_run("\t \t \t \tIncident ref: ")
    font = titre1.font
    font.name = 'Calibri (Body)'
    font.size = Pt(16)

    titre2 = parag.add_run(type+"-"+date+"-"+numero+"\n")
    font = titre2.font
    font.name = 'Calibri (Body)'
    font.size = Pt(16)
    font.bold = True

    if (serv == "Elévé"):
        font.color.rgb = RGBColor(255, 0, 0)
    elif (serv == "Moyenne"):
        font.color.rgb = RGBColor(255, 128, 0)
    else:
        font.color.rgb = RGBColor(0, 153, 0)

    parag.add_run("\n")
    titre3 = parag.add_run("Sujet: "+sujet+"\n")
    if " " in sujet:
        sujet = sujet.replace(" ","_")
    font = titre3.font
    font.name = 'Calibri (Body)'
    font.size = Pt(14)
    font.bold = True
    font.color.rgb = RGBColor(255, 255, 255)
    if (serv == "Elévé"):
        font.highlight_color = WD_COLOR_INDEX.RED
    elif (serv == "Moyenne"):
        font.highlight_color = WD_COLOR_INDEX.YELLOW
    else:
        font.highlight_color = WD_COLOR_INDEX.GREEN

    titre4 = parag.add_run("Réference: " )
    font = titre4.font
    font.name = 'Calibri (Body)'
    font.size = Pt(14)
    font.bold = True

    titre5 = parag.add_run(type + "-" + date + "-" + numero+"\n")
    font = titre5.font
    font.name = 'Calibri (Body)'
    font.size = Pt(16)
    font.bold = True

    if (serv == "Elévé"):
        font.color.rgb = RGBColor(255, 0, 0)
    elif (serv == "Moyenne"):
        font.color.rgb = RGBColor(255, 128, 0)
    else:
        font.color.rgb = RGBColor(0, 153, 0)

    titre6 = parag.add_run("Sévirité: ")
    font = titre6.font
    font.name = 'Calibri (Body)'
    font.size = Pt(14)
    font.bold = True

    titre7 = parag.add_run(serv+"\n")
    font = titre7.font
    font.name = 'Calibri (Body)'
    font.size = Pt(14)
    font.bold = True

    if (serv == "Elévé"):
        font.color.rgb = RGBColor(255, 0, 0)
    elif (serv == "Moyenne"):
        font.color.rgb = RGBColor(255, 128, 0)
    else:
        font.color.rgb = RGBColor(0, 153, 0)

    titre8 = parag.add_run("Description:\n")
    font = titre8.font
    font.name = 'Calibri (Body)'
    font.size = Pt(14)
    font.bold = True

    titre9 = parag.add_run(description+"\n")
    font = titre9.font
    font.name = 'Calibri (Body)'
    font.size = Pt(12)

    titre99 = parag.add_run("Log Message :\n")
    font = titre99.font
    font.name = 'Calibri (Body)'
    font.size = Pt(14)
    font.bold = True

    titre10 = parag.add_run(logmesg + "\n")
    font = titre10.font
    font.name = 'Times New Roman'
    font.size = Pt(9)

    titre11 = parag.add_run("Source :\n")
    font = titre11.font
    font.name = 'Calibri (Body)'
    font.size = Pt(14)
    font.bold = True

    titre12 = parag.add_run("Logrhythm \n")
    font = titre12.font
    font.name = 'Calibri (Body)'
    font.size = Pt(12)

    titre13 = parag.add_run("Recommandation :\n")
    font = titre13.font
    font.name = 'Calibri (Body)'
    font.size = Pt(14)
    font.bold = True

    titre14 = parag.add_run(recommandation)
    font = titre14.font
    font.name = 'Calibri (Body)'
    font.size = Pt(12)

    doc.save("ATB-" + type + "-" + date + "-" + "0" + numero + "_" + sujet + ".docx")
    os.system("start ATB-" + type + "-" + date + "-" + "0" + numero + "_" + sujet + ".docx")
    return render_template('/rapport.html')

@app.route('/wazuh')
def wazuh1():
    return render_template('/wazuh.html')


@app.route('/wazuh', methods=['get', 'post'])
def wazuh():
    client = request.form['client']
    sujet = request.form['sujet']
    type = request.form['type']
    serv = request.form['serv']
    numero = request.form['numero']
    img = request.form['description']
    description = request.form['description']
    recommandation = request.form['recommandation']

    requete_wazuh = """INSERT INTO wazuh(type,numero,sujet,date,client,description,recommandation) VALUES(%s, %s, %s, %s, %s, %s)"""
    recordTuple = (type, numero, sujet, date, client, description, recommandation)
    myCursor.execute(requete_wazuh, recordTuple)
    mydb.commit()


    doc = DocxTemplate("ATB_Template.docx")
    context = {'Keystone': "World company"}
    doc.render(context)
    parag = doc.add_paragraph()

    titre1 = parag.add_run("\t \t \t \tIncident ref: ")
    font = titre1.font
    font.name = 'Calibri (Body)'
    font.size = Pt(16)

    titre2 = parag.add_run(type + "-" + date + "-" + numero + "\n")
    font = titre2.font
    font.name = 'Calibri (Body)'
    font.size = Pt(16)
    font.bold = True

    if (serv == "Elévé"):
        font.color.rgb = RGBColor(255, 0, 0)
    elif (serv == "Moyenne"):
        font.color.rgb = RGBColor(255, 128, 0)
    else:
        font.color.rgb = RGBColor(0, 153, 0)

    parag.add_run("\n")
    titre3 = parag.add_run("Sujet: " + sujet + "\n")
    if " " in sujet:
        sujet = sujet.replace(" ", "_")
    font = titre3.font
    font.name = 'Calibri (Body)'
    font.size = Pt(14)
    font.bold = True
    font.color.rgb = RGBColor(255, 255, 255)
    if (serv == "Elévé"):
        font.highlight_color = WD_COLOR_INDEX.RED
    elif (serv == "Moyenne"):
        font.highlight_color = WD_COLOR_INDEX.YELLOW
    else:
        font.highlight_color = WD_COLOR_INDEX.GREEN

    titre4 = parag.add_run("Réference: ")
    font = titre4.font
    font.name = 'Calibri (Body)'
    font.size = Pt(14)
    font.bold = True

    titre5 = parag.add_run(type + "-" + date + "-" + numero + "\n")
    font = titre5.font
    font.name = 'Calibri (Body)'
    font.size = Pt(16)
    font.bold = True

    if (serv == "Elévé"):
        font.color.rgb = RGBColor(255, 0, 0)
    elif (serv == "Moyenne"):
        font.color.rgb = RGBColor(255, 128, 0)
    else:
        font.color.rgb = RGBColor(0, 153, 0)

    titre6 = parag.add_run("Sévirité: ")
    font = titre6.font
    font.name = 'Calibri (Body)'
    font.size = Pt(14)
    font.bold = True

    titre7 = parag.add_run(serv + "\n")
    font = titre7.font
    font.name = 'Calibri (Body)'
    font.size = Pt(14)
    font.bold = True

    if (serv == "Elévé"):
        font.color.rgb = RGBColor(255, 0, 0)
    elif (serv == "Moyenne"):
        font.color.rgb = RGBColor(255, 128, 0)
    else:
        font.color.rgb = RGBColor(0, 153, 0)

    titre8 = parag.add_run("Description:\n")
    font = titre8.font
    font.name = 'Calibri (Body)'
    font.size = Pt(14)
    font.bold = True

    titre9 = parag.add_run(description + "\n")
    font = titre9.font
    font.name = 'Calibri (Body)'
    font.size = Pt(12)

    titre11 = parag.add_run("Source :\n")
    font = titre11.font
    font.name = 'Calibri (Body)'
    font.size = Pt(14)
    font.bold = True

    titre12 = parag.add_run("Wazuh \n")
    font = titre12.font
    font.name = 'Calibri (Body)'
    font.size = Pt(12)

    titre13 = parag.add_run("Recommandation :\n")
    font = titre13.font
    font.name = 'Calibri (Body)'
    font.size = Pt(14)
    font.bold = True

    titre14 = parag.add_run(recommandation)
    font = titre14.font
    font.name = 'Calibri (Body)'
    font.size = Pt(12)

    doc.save(client + "-" + type + "-" + date + "-" + "0" + numero + "_" + sujet + ".docx")
    os.system("start " + client + "-" + type + "-" + date + "-" + "0" + numero + "_" + sujet + ".docx")
    return render_template('/wazuh.html')